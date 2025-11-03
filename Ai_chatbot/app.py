import google.generativeai as ai
import pdfplumber
import os
import ollama
from PIL import Image
import pytesseract
from flask import Flask, render_template, request, jsonify, send_from_directory, session, request, redirect
import uuid
from dotenv import load_dotenv
import docx
import boto3
import google.generativeai as genai
import cv2
import numpy as np
import pandas as pd
from tabulate import tabulate
import re
from datetime import datetime
from email_validator import clean_email, is_valid_email_format, extract_email_addresses_improved
from document_analyzer import DocumentAnalyzer
from auth import verify_user, get_user_email, extract_bearer_token
from db import add_document, list_user_documents, create_chat, add_message, list_user_chats, list_messages, upsert_profile, list_profiles, upsert_profile_as_user, get_user_messages, append_user_message, get_user_messages_as_user, append_user_message_as_user
from storage import upload_file_to_b2
import tempfile
import requests
import threading
from PyPDF2 import PdfReader
import json

# Load environment variables
load_dotenv()

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size
app.secret_key = os.getenv('SECRET_KEY', 'add8e6b193e17b39dd36ab85216f2c4c')

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
LOG_DIR = os.path.join(BASE_DIR, "logs")
CHAT_LOG_PATH = os.path.join(LOG_DIR, "chat_log.txt")
PERSIST_DIR = os.path.join(BASE_DIR, "persist")
CLIENT_MAP_PATH = os.path.join(PERSIST_DIR, "client_uploads.json")

os.makedirs(PERSIST_DIR, exist_ok=True)

def log_conversation(user_input, bot_response, override=False):
    try:
        log_path = session.get('chat_log_path')
        if not log_path:
            print("⚠️ No log path set for this session.")
            return
        full_path = os.path.join(os.getcwd(), log_path)
        os.makedirs(os.path.dirname(full_path), exist_ok=True)
        mode = 'w' if override else 'a'
        with open(full_path, mode, encoding='utf-8') as f:
            f.write(f"User: {user_input}\n")
            f.write(f"Bot: {bot_response}\n\n")
        print("📝 Logged to:", full_path)
    except Exception as e:
        print("❌ Logging error:",e)


        os.makedirs(LOG_DIR, exist_ok=True)
        mode = 'w' if override else 'a'
        with open(CHAT_LOG_PATH, mode, encoding='utf-8') as f:
            f.write(f"User: {user_input}\n")
            f.write(f"Bot: {bot_response}\n\n")
        print("📝 Chat logged successfully at ", CHAT_LOG_PATH)
    except Exception as e:
        print("❌ Error logging chat:", e)

ALLOWED_EXTENSIONS = {'pdf', 'txt', 'doc', 'docx', 'rtf'}
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def generate_polly_audio(text, filename="response.mp3", voice="Matthew"):

    static_path=os.path.join('Ai_chatbot', 'static')
    os.makedirs(static_path, exist_ok=True)

    output_path=os.path.join(static_path,filename)
    try :
        # Check AWS credentials before creating client to avoid runtime noise
        has_access_key = os.getenv('AWS_ACCESS_KEY_ID') and os.getenv('AWS_SECRET_ACCESS_KEY')
        if not has_access_key:
            print("[Polly ERROR] AWS credentials not configured; skipping TTS")
            return None
        polly=boto3.client('polly', region_name='us-east-1')

        response=polly.synthesize_speech(
            Text=text,
            OutputFormat='mp3',
            VoiceId=voice
        )

        print("[Polly RAW RESPONSE]", response)
        if 'AudioStream' not in response:
            print("[POLLY ERROR] AudioStream not in response")
            if 'error' in response:
                print("[Polly ERROR DETAILS]:", response['Error'])
        
        with open(output_path,'wb') as f:
            f.write(response['AudioStream'].read())
        return f"/static/{filename}"
    
    except Exception as e:
        print("[Polly ERROR]", str(e))
        return None


def load_client_map():
    try:
        if os.path.exists(CLIENT_MAP_PATH):
            with open(CLIENT_MAP_PATH, "r", encoding="utf-8") as f:
                return json.load(f)
    except Exception as e:
        print("[PERSIST] Failed to load client map:", e)
    return {}


def save_client_map(mapping):
    try:
        with open(CLIENT_MAP_PATH, "w", encoding="utf-8") as f:
            json.dump(mapping, f, ensure_ascii=False, indent=2)
    except Exception as e:
        print("[PERSIST] Failed to save client map:", e)


def get_or_create_client_id():
    try:
        cid = session.get('client_id')
        if not cid:
            cid = str(uuid.uuid4())
            session['client_id'] = cid
        return cid
    except Exception as e:
        print("[SESSION] Failed to get/create client_id:", e)
        return str(uuid.uuid4())


def extract_text_from_pdf(path):
        text = ""
        with open(path, 'rb') as f:
            reader = PdfReader(f)
            for page in reader.pages:
                text += page.extract_text() or ''
        return text

def _is_http_url(p):
        return isinstance(p, str) and (p.startswith('http://') or p.startswith('https://'))

def ensure_local_file(source_path_or_url):
        """If given a URL, download to a temp file and return the local path; otherwise return path as-is."""
        try:
            if _is_http_url(source_path_or_url):
                resp = requests.get(source_path_or_url, stream=True, timeout=30)
                resp.raise_for_status()
                suffix = os.path.splitext(source_path_or_url.split('?')[0])[1] or '.bin'
                fd, tmp_path = tempfile.mkstemp(suffix=suffix)
                with os.fdopen(fd, 'wb') as out:
                    for chunk in resp.iter_content(chunk_size=8192):
                        if chunk:
                            out.write(chunk)
                return tmp_path
        except Exception as e:
            print('[DOWNLOAD] Failed to fetch', source_path_or_url, e)
        return source_path_or_url

def get_b2_client():
        try:
            endpoint = os.getenv('B2_ENDPOINT')
            key_id = os.getenv('B2_KEY_ID')
            app_key = os.getenv('B2_APP_KEY')
            if not (endpoint and key_id and app_key):
                return None
            return boto3.client('s3', endpoint_url=endpoint, aws_access_key_id=key_id, aws_secret_access_key=app_key)
        except Exception as e:
            print('[B2] client init failed:', e)
            return None

def upload_to_b2(local_path, object_key):
        s3 = get_b2_client()
        if not s3:
            return None
        bucket = os.getenv('B2_BUCKET')
        public_base = os.getenv('B2_PUBLIC_URL') or os.getenv('B2_ENDPOINT')
        if not bucket:
            return None
        try:
            s3.upload_file(local_path, bucket, object_key)
            if public_base:
                base = public_base.rstrip('/')
                return f"{base}/{bucket}/{object_key}"
            return None
        except Exception as e:
            print('[B2] upload failed:', e)
            return None

# Create uploads folder if it doesn't exist
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

class EnhancedMultiFormatChatbot:
    def __init__(self, api_key):
        ai.configure(api_key=api_key)
        # Discover available models for this API key and pick those supporting generateContent
        self.model_candidates = self._discover_models()
        # Initialize with the first candidate; will fallback on first use
        initial_model_name = self.model_candidates[0] if self.model_candidates else "gemini-1.5-flash"
        print(f"[Gemini] Initial model candidate: {initial_model_name}")
        self.model = ai.GenerativeModel(initial_model_name)
        self.content = ""
        self.chat = None
        self.current_file = ""
        self.extracted_tables = []
        self.document_analyzer = DocumentAnalyzer()

    def _discover_models(self):
        """Return an ordered list of model IDs available to this key that support generateContent."""
        candidates = []
        try:
            models = ai.list_models()
            # Prefer 1.5 flash, then 1.5 pro, then others
            preferred_order = [
                "gemini-1.5-flash",
                "gemini-1.5-pro",
                "gemini-1.0-pro",
                "gemini-pro",
            ]
            # Build map from base to full name
            supported = []
            for m in models:
                methods = getattr(m, "supported_generation_methods", []) or []
                if "generateContent" in methods:
                    supported.append(m.name)  # full name like 'models/gemini-1.5-flash'
            print(f"[Gemini] Supported models for this key: {supported}")
            # Order supported by our preference if base name matches
            def base(n):
                return n.split("/")[-1]
            # Preferred
            for p in preferred_order:
                for full in supported:
                    if base(full) == p:
                        candidates.append(full)
                        break
            # Add any remaining supported not already included
            for full in supported:
                if full not in candidates:
                    candidates.append(full)
        except Exception as e:
            print(f"[Gemini] Failed to list models, falling back to defaults: {e}")
            candidates = [
                "gemini-1.5-flash",
                "gemini-1.5-pro",
                "gemini-pro",
            ]
        return candidates

    def _start_chat_with_prompt(self, prompt):
        """Try starting a chat with multiple model candidates until one works."""
        last_error = None
        for model_name in self.model_candidates:
            try:
                self.model = ai.GenerativeModel(model_name)
                self.chat = self.model.start_chat()
                _ = self.chat.send_message(prompt)
                print(f"✅ Chat initialized with model: {model_name}")
                return True
            except Exception as e:
                print(f"[WARN] Model '{model_name}' failed: {e}")
                self.chat = None
                last_error = e
        # As a last resort, try the short base ids
        for model_name in ["gemini-1.5-flash", "gemini-1.5-pro", "gemini-pro"]:
            try:
                self.model = ai.GenerativeModel(model_name)
                self.chat = self.model.start_chat()
                _ = self.chat.send_message(prompt)
                print(f"✅ Chat initialized with fallback model: {model_name}")
                return True
            except Exception as e:
                print(f"[WARN] Fallback model '{model_name}' failed: {e}")
                self.chat = None
                last_error = e
        print(f"[ERROR] All model candidates failed. Last error: {last_error}")
        return False

    def set_context(self, text):
        self.context = text

    def load_combined_text(self, combined_text):
        """Load multiple PDF contents together into a single chat context"""
        self.text_context = combined_text

        context_prompt = f"""
        You are an AI assistant that answers questions based on document content, including structured data like tables and forms.

    IMPORTANT INSTRUCTIONS:
    - Give direct, accurate answers based on the extracted content
    - Pay special attention to tabular data and structured information
    - When asked about educational qualifications, look for institution names, years, and percentages
    - For B.Tech questions, specifically look for university/college names in the education section
    - Be specific with names, dates, and numbers when available
    - If information is in a table format, extract the exact values

    DOCUMENT CONTENT:
        {self.text_context}

    The document has been processed with enhanced table extraction. Answer questions accurately based on this content.
    """

        try:
            if self._start_chat_with_prompt(context_prompt):
                print("✅ Combined documents loaded successfully!")
                return True
        except Exception as e:
            print(f"[ERROR] Failed to load combined content into chat: {e}")
            self.chat = None
        return False

    def preprocess_image_for_ocr(self, image_path):
        """Preprocess image to improve OCR accuracy"""
        try:
            # Read image using OpenCV
            img = cv2.imread(image_path)
            
            # Convert to grayscale
            gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
            
            # Apply noise reduction
            denoised = cv2.fastNlMeansDenoising(gray)
            
            # Apply adaptive thresholding to get better contrast
            thresh = cv2.adaptiveThreshold(
                denoised, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 11, 2
            )
            
            # Morphological operations to clean up the image
            kernel = np.ones((2, 2), np.uint8)
            processed = cv2.morphologyEx(thresh, cv2.MORPH_CLOSE, kernel)
            
            return processed
            
        except Exception as e:
            print(f"Error in image preprocessing: {e}")
            return cv2.imread(image_path, cv2.IMREAD_GRAYSCALE)
    
    def detect_and_extract_tables(self, image_path):
        """Enhanced table detection and extraction"""
        try:
            img = cv2.imread(image_path)
            gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
            
            # Apply threshold to get binary image
            thresh = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY_INV + cv2.THRESH_OTSU)[1]
            
            # Detect horizontal lines
            horizontal_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (25, 1))
            horizontal_lines = cv2.morphologyEx(thresh, cv2.MORPH_OPEN, horizontal_kernel, iterations=2)
            
            # Detect vertical lines
            vertical_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (1, 25))
            vertical_lines = cv2.morphologyEx(thresh, cv2.MORPH_OPEN, vertical_kernel, iterations=2)
            
            # Combine horizontal and vertical lines
            table_mask = cv2.addWeighted(horizontal_lines, 0.5, vertical_lines, 0.5, 0.0)
            
            # Find contours to detect table cells
            contours, _ = cv2.findContours(table_mask, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
            
            # Sort contours by area and keep only significant ones
            contours = sorted(contours, key=cv2.contourArea, reverse=True)
            significant_contours = [c for c in contours if cv2.contourArea(c) > 100]
            
            return len(significant_contours) > 3, significant_contours
            
        except Exception as e:
            print(f"Error in table detection: {e}")
            return False, []
    
    def extract_table_with_coordinates(self, image_path):
        """Extract table data using coordinate-based approach"""
        try:
            image = Image.open(image_path)
            
            # Get OCR data with coordinates
            ocr_data = pytesseract.image_to_data(
                image, 
                config='--psm 6 -c tessedit_char_whitelist=ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz0123456789.,()-%/ ',
                output_type=pytesseract.Output.DICT
            )
            
            # Filter confident detections
            confident_data = []
            for i in range(len(ocr_data['text'])):
                if int(ocr_data['conf'][i]) > 40:  # Lower threshold for better detection
                    text = ocr_data['text'][i].strip()
                    if text and len(text) > 1:  # Ignore single characters
                        confident_data.append({
                            'text': text,
                            'left': ocr_data['left'][i],
                            'top': ocr_data['top'][i],
                            'width': ocr_data['width'][i],
                            'height': ocr_data['height'][i]
                        })
            
            if not confident_data:
                return None
            
            # Group data into rows based on Y coordinates
            rows = {}
            y_tolerance = 15  # Pixels tolerance for same row
            
            for item in confident_data:
                y = item['top']
                
                # Find existing row or create new one
                row_key = None
                for existing_y in rows.keys():
                    if abs(y - existing_y) <= y_tolerance:
                        row_key = existing_y
                        break
                
                if row_key is None:
                    row_key = y
                    rows[row_key] = []
                
                rows[row_key].append(item)
            
            # Sort rows by Y coordinate and create table
            sorted_rows = sorted(rows.keys())
            table_data = []
            
            for row_y in sorted_rows:
                # Sort items in row by X coordinate
                row_items = sorted(rows[row_y], key=lambda x: x['left'])
                row_text = [item['text'] for item in row_items]
                
                if len(row_text) > 1:  # Only include rows with multiple columns
                    table_data.append(row_text)
            
            return table_data
            
        except Exception as e:
            print(f"Error in coordinate-based extraction: {e}")
            return None
    
    def format_extracted_table(self, table_data):
        """Format extracted table data into readable text"""
        if not table_data:
            return ""
        
        try:
            # Create DataFrame
            if len(table_data) > 1:
                # First row as headers if it looks like headers
                headers = table_data[0]
                data_rows = table_data[1:]
                
                # Pad rows to match header length
                max_cols = len(headers)
                for i, row in enumerate(data_rows):
                    while len(row) < max_cols:
                        row.append("")
                    data_rows[i] = row[:max_cols]  # Trim if too long
                
                df = pd.DataFrame(data_rows, columns=headers)
                formatted_table = df.to_string(index=False)
                
                # Also create a structured text representation
                structured_text = f"\n=== TABLE STRUCTURE ===\n"
                for _, row in df.iterrows():
                    for col, val in row.items():
                        if val.strip():
                            structured_text += f"{col}: {val}\n"
                    structured_text += "---\n"
                
                return f"{formatted_table}\n{structured_text}"
            else:
                return str(table_data)
                
        except Exception as e:
            print(f"Error formatting table: {e}")
            return str(table_data)
    
    def extract_text_with_enhanced_table_support(self, image_path):
        """Enhanced text extraction with better table and section support"""
        try:
            print("Starting enhanced image extraction with section analysis...")
            
            image = Image.open(image_path)
            if image.mode != 'RGB':
                image = image.convert('RGB')
                
            # Get detailed OCR data with bounding boxes
            ocr_data = pytesseract.image_to_data(image, output_type=pytesseract.Output.DICT, config='--psm 6')
            
            # Process text elements
            text_elements = []
            for i in range(len(ocr_data['text'])):
                text = ocr_data['text'][i].strip()
                conf = int(ocr_data['conf'][i])
                if text and conf > 50:
                     text_elements.append({
                         'text': text,
                         'left': ocr_data['left'][i],
                         'top': ocr_data['top'][i],
                         'width': ocr_data['width'][i],
                         'height': ocr_data['height'][i],
                         'line_num': ocr_data['line_num'][i],
                         'block_num': ocr_data['block_num'][i]
                     })
            
            # Sort elements by position
            text_elements.sort(key=lambda x: (x['top'], x['left']))
            
            # Reconstruct lines
            lines = []
            current_line = []
            y_tolerance = 10
            
            for element in text_elements:
                if not current_line:
                    current_line.append(element)
                else:
                    avg_line_height = sum(el['height'] for el in current_line) / len(current_line)
                    if abs(element['top'] - current_line[-1]['top']) < y_tolerance or \
                       (element['top'] + element['height'] > current_line[-1]['top'] and element['top'] < current_line[-1]['top'] + current_line[-1]['height']):
                         current_line.append(element)
                    else:
                        lines.append(current_line)
                        current_line = [element]
            
            if current_line:
                lines.append(current_line)
                
            # Process lines
            processed_lines = []
            for line_elements in lines:
                 line_elements.sort(key=lambda x: x['left'])
                 processed_lines.append(" ".join([el['text'] for el in line_elements]))

            # Analyze document content
            text_content = "\n".join(processed_lines)
            analysis = self.document_analyzer.analyze_document(text_content, image_path)
            
            # Format the analysis results
            structured_text = self.document_analyzer.format_analysis(analysis)
            
            # Add raw OCR outputs for redundancy
            all_ocr_outputs = ""
            all_ocr_outputs += "\n=== RAW OCR (PSM 6) ===\n" + pytesseract.image_to_string(image, config='--psm 6') + "\n\n"
            all_ocr_outputs += "=== RAW OCR (PSM 4) ===\n" + pytesseract.image_to_string(image, config='--psm 4') + "\n\n"
            all_ocr_outputs += "=== RAW OCR (PSM 11) ===\n" + pytesseract.image_to_string(image, config='--psm 11') + "\n\n"

            structured_text += all_ocr_outputs
            
            return structured_text.strip()
            
        except Exception as e:
            print(f"Error in enhanced image extraction: {e}")
            return self.extract_text_from_image_standard(image_path)
    
    def extract_email_addresses(self, text):
        """Extract email addresses from text using improved extraction"""
        return extract_email_addresses_improved(text, self.current_image_path if hasattr(self, 'current_image_path') else None)
    
    def extract_education_patterns(self, text):
        """Extract education-related patterns from text, including percentages"""
        try:
            education_info = []
            
            # Look for B.Tech patterns
            btech_patterns = [
                r'B\.?Tech\.?\s*(?:\([^)]*\))?\s*[:\-]?\s*([A-Za-z\s&]+(?:College|University|Institute))',
                r'B\.?Tech\.?\s*(?:\([^)]*\))?\s*([A-Za-z\s&]+(?:College|University|Institute))',
                r'(?:B\.?Tech\.?.*?)([A-Za-z\s&]+(?:Engineering|College|University|Institute))',
            ]
            
            for pattern in btech_patterns:
                matches = re.finditer(pattern, text, re.IGNORECASE)
                for match in matches:
                    institution = match.group(1).strip()
                    if len(institution) > 3:  # Filter out too short matches
                        education_info.append(f"B.Tech Institution: {institution}")
            
            # Look for year patterns near B.Tech
            year_patterns = r'B\.?Tech\.?.*?(?:20\d{2}|19\d{2})'
            year_matches = re.finditer(year_patterns, text, re.IGNORECASE)
            for match in year_matches:
                education_info.append(f"B.Tech Year Info: {match.group()}")
            
            # Look for percentage patterns specifically near education keywords (refined)
            percentage_patterns_education = r'(?:B\.?Tech\.?|Diploma|H\.?S\.?C\.?E\.?|S\.?S\.?C\.?E\.?).*?(\d{1,2}(?:.\d{1,2})?)\s*%?' # Look for 1-2 digits, optional decimal, optional %
            percentage_matches_education = re.finditer(percentage_patterns_education, text, re.IGNORECASE)
            for match in percentage_matches_education:
                percentage = match.group(1).strip()
                if percentage:
                    # Basic validation: ensure percentage is reasonable (e.g., <= 100)
                    try:
                        percentage_value = float(percentage)
                        if 0 <= percentage_value <= 100:
                            education_info.append(f"Education Percentage Found (Refined): {percentage}%")
                        else:
                            print(f"Warning: Extracted percentage {percentage}% seems unreasonable.")
                    except ValueError:
                        print(f"Warning: Could not convert extracted percentage {percentage} to float.")
            
            # Broader percentage pattern as a fallback (also refined)
            percentage_patterns_general = r'\b(\d{1,2}(?:.\d{1,2})?)\s*%?' # Look for 1-2 digits, optional decimal, optional %
            percentage_matches_general = re.finditer(percentage_patterns_general, text)
            for match in percentage_matches_general:
                percentage = match.group(1).strip()
                if percentage:
                    # Basic validation: ensure percentage is reasonable (e.g., <= 100)
                    try:
                        percentage_value = float(percentage)
                        if 0 <= percentage_value <= 100 and f"Education Percentage Found (Refined): {percentage}%" not in education_info:
                            education_info.append(f"Percentage Found (General, Refined): {percentage}%")
                    except ValueError:
                        print(f"Warning: Could not convert extracted percentage {percentage} to float (General).")
            
            return '\n'.join(education_info) if education_info else ""
            
        except Exception as e:
            print(f"Error in pattern extraction: {e}")
            return ""
    
    def extract_text_from_pdf(self, pdf_path):
        """Extract text using pdfplumber for PDF files with table support"""
        try:
            text = ""
            tables_text = ""
            
            with pdfplumber.open(pdf_path) as pdf:
                print(f"Processing {len(pdf.pages)} pages...")
                
                for page_num, page in enumerate(pdf.pages, 1):
                    # Extract regular text
                    page_text = page.extract_text()
                    if page_text:
                        text += f"\n--- Page {page_num} ---\n{page_text}"
                    
                    # Extract tables
                    tables = page.extract_tables()
                    if tables:
                        tables_text += f"\n--- Tables from Page {page_num} ---\n"
                        for table_num, table in enumerate(tables, 1):
                            tables_text += f"\nTable {table_num}:\n"
                            # Convert table to readable format
                            if table:
                                df = pd.DataFrame(table[1:], columns=table[0] if table[0] else None)
                                tables_text += df.to_string(index=False) + "\n"
                    
                    print(f"Processed page {page_num}")
            
            # Combine regular text and tables
            combined_text = text
            if tables_text:
                combined_text += "\n\n=== EXTRACTED TABLES ===\n" + tables_text
                
            return combined_text.strip()
        
        except Exception as e:
            print(f"Error reading PDF: {e}")
            return None
    
    def extract_text_from_image_standard(self, image_path):
        """Standard image text extraction"""
        try:
            print("Using standard OCR extraction...")
            
            image = Image.open(image_path)
            if image.mode != 'RGB':
                image = image.convert('RGB')
            
            extracted_text = pytesseract.image_to_string(image, config='--psm 6')
            
            if extracted_text.strip():
                return extracted_text.strip()
            else:
                return "No readable text found in the image."
                
        except Exception as e:
            print(f"Error reading image: {e}")
            return None
    
    def extract_text_from_image(self, image_path):
        """Main image text extraction method"""
        return self.extract_text_with_enhanced_table_support(image_path)
    
    def extract_text_from_docx(self, docx_path):
        """Extract text from DOC/DOCX files"""
        try:
            text = ""
            doc = docx.Document(docx_path)
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract tables from docx
            tables_text = ""
            for table_num, table in enumerate(doc.tables, 1):
                tables_text += f"\nTable {table_num}:\n"
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                
                if table_data:
                    df = pd.DataFrame(table_data[1:], columns=table_data[0] if table_data[0] else None)
                    tables_text += df.to_string(index=False) + "\n"
            
            combined_text = text
            if tables_text:
                combined_text += "\n\n=== EXTRACTED TABLES ===\n" + tables_text
                
            return combined_text.strip()
        except Exception as e:
            print(f"Error reading DOCX: {e}")
            return None
    
    def load_file(self, file_path):
        """Load PDF, DOC/DOCX or image file and prepare for Q&A"""
        if not os.path.exists(file_path):
            print(f"Error: File '{file_path}' not found!")
            return False
        
        # Determine file type
       
        file_extension = os.path.splitext(file_path.lower())[1]
        
        print(f"Loading file: {os.path.basename(file_path)}")
        
        if file_extension == '.pdf':
            self.content = self.extract_text_from_pdf(file_path)
            file_type = "PDF"
        elif file_extension in ['.doc', '.docx']:
            self.content = self.extract_text_from_docx(file_path)
            file_type = "DOC/DOCX"
        elif file_extension in ['.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.webp']:
            self.content = self.extract_text_from_image(file_path)
            file_type = "Image"
        else:
            print(f"Unsupported file format: {file_extension}")
            print("Supported formats: PDF, DOC, DOCX, JPG, JPEG, PNG, BMP, TIFF, WEBP")
            return False
        
        if not self.content:
            print(f"Failed to extract text from {file_type}!")
            return False
        
        # Show preview of extracted content for debugging
        print("\n=== EXTRACTED CONTENT PREVIEW ===")
        print(self.content[:1000] + "..." if len(self.content) > 1000 else self.content)
        print("=== END PREVIEW ===\n")
        
        # Truncate content if too long
        if len(self.content) > 50000:
            self.content = self.content[:50000] + "\n[Content truncated due to length...]"
            print("Note: Content was truncated due to length limits.")
        
        self.current_file = os.path.basename(file_path)
        
        # Enhanced context prompt for better table understanding
        context_prompt = f"""
        You are an AI assistant that answers questions based on document content, including structured data like tables and forms.
        
        IMPORTANT INSTRUCTIONS:
        - Give direct, accurate answers based on the extracted content
        - Pay special attention to tabular data and structured information
        - When asked about educational qualifications, look for institution names, years, and percentages
        - For B.Tech questions, specifically look for university/college names in the education section
        - Be specific with names, dates, and numbers when available
        - If information is in a table format, extract the exact values
        
        DOCUMENT CONTENT:
        {self.content}
        
        The document has been processed with enhanced table extraction. Answer questions accurately based on this content.
        """
        
        try:
            if self._start_chat_with_prompt(context_prompt):
                print(f"✅ {file_type} loaded successfully! Enhanced table extraction completed.")
                return True
        except Exception as e:
            print(f"Error initializing chat: {e}")
        return False
    
    def ask_question(self, question):
        """Ask a question with enhanced context understanding"""
        if not self.chat:
            return "Please load a file first using 'load <file_path>' command."
        
        try:
            # Check if the question is about a specific year
            specific_year_match = re.search(r'(20\d{2})', question.lower())
            is_general_query = not specific_year_match

            # Prepare combined context
            base_text_for_filter = getattr(self, 'text_context', None) or self.content or ""

            # Import helpers
            from content_filter import filter_latest_year_content, get_latest_year_prompt

            # Always compute latest-year filtered content (used by both general and specific queries)
            filtered_content, latest_year = filter_latest_year_content(base_text_for_filter)
            if is_general_query and latest_year:
                # Modify question to steer the model
                question = f"{question} for {latest_year} ONLY"
                print(f"[FILTER] Modified question to: {question}")
            # Cache filtered content for extraction and prompts
            self.filtered_content = filtered_content or ""

            # If the user asked for a specific leave type count, answer concisely without the model
            try:
                from content_filter import extract_leave_days, detect_leave_type_from_question
                ql = question.lower()
                detected_type = detect_leave_type_from_question(ql)
                if detected_type:
                    # Prefer latest-year filtered content, fallback to combined text
                    days = extract_leave_days(self.filtered_content, detected_type) or extract_leave_days(base_text_for_filter, detected_type)
                    if days:
                        return days
            except Exception as _e:
                pass
            
            # Enhanced question processing with year-specific instructions and content filtering
            if is_general_query and 'latest_year' in locals() and latest_year:
                # Determine the base text to operate on (combined docs take precedence)
                base_text = getattr(self, 'text_context', None) or self.content or ""
                # Prefer the strict filtered content if available; otherwise use the base text
                effective_text = getattr(self, 'filtered_content', None) or base_text
                # Use our dedicated function to get a restrictive system prompt
                base_question_prompt = get_latest_year_prompt(question, latest_year)
            else:
                base_question_prompt = f"""
Question: {question}

Please provide a direct, specific answer based on the document content. The document has been analyzed for:
- Education details (degrees, institutions, years, scores)
- Contact information (email, phone, address)
- Skills and qualifications
- Professional experience
- Projects and achievements
- Other relevant sections

Please focus on the specific information requested in the question.
"""

            # Construct the full prompt with the correct content source
            if is_general_query and 'latest_year' in locals() and latest_year:
                # Use the filtered combined text if available, to avoid prior chat context
                content_source = (effective_text if 'effective_text' in locals() and effective_text else (getattr(self, 'text_context', None) or self.content))
                full_prompt = base_question_prompt + "\n\nDocument Content (filtered):\n" + content_source
                # Send via a fresh, ephemeral chat to avoid contamination from earlier messages
                temp_chat = self.model.start_chat()
                response = temp_chat.send_message(full_prompt)
                response_text = response.text if hasattr(response, 'text') else str(response)
                # Post-process: strip any lines mentioning non-latest years
                try:
                    lines = response_text.splitlines()
                    cleaned_lines = []
                    for ln in lines:
                        years_in_line = set(re.findall(r'20\d{2}', ln))
                        if any(y != latest_year for y in years_in_line):
                            continue
                        cleaned_lines.append(ln)
                    response_text_clean = "\n".join(cleaned_lines).strip()
                    # If this was a specific leave-type query, reduce to concise days
                    try:
                        from content_filter import extract_leave_days, detect_leave_type_from_question
                        detected_type = detect_leave_type_from_question(question)
                        if detected_type:
                            concise = extract_leave_days(response_text_clean or response_text, detected_type)
                            if concise:
                                return concise
                    except Exception:
                        pass
                    if not response_text_clean:
                        response_text_clean = f"Here is the information for {latest_year} (latest available):\nNo other year-specific details are included."
                    return response_text_clean
                except Exception:
                    return response_text
            else:
                full_prompt = base_question_prompt + "\n\nDocument Content:\n" + self.content
                response = self.chat.send_message(full_prompt)
                response_text = response.text if hasattr(response, 'text') else str(response)
                # If the question was a specific leave-type query but extraction failed earlier,
                # try to reduce the model's verbose answer to just the days.
                try:
                    from content_filter import extract_leave_days, detect_leave_type_from_question
                    detected_type = detect_leave_type_from_question(question)
                    if detected_type:
                        concise = extract_leave_days(response_text, detected_type)
                        if concise:
                            return concise
                except Exception:
                    pass
                return response_text

        except Exception as e:
            return f"Error: {e}"
    
    def show_content_preview(self):
        """Show a preview of the extracted content"""
        if not self.content:
            print("No content loaded.")
            return
        
        preview = self.content[:1000] + "..." if len(self.content) > 1000 else self.content
        print(f"\n--- Content Preview from {self.current_file} ---")
        print(preview)
        print("--- End Preview ---\n")

# Initialize chatbot with API key from environment variable
API_KEY = os.getenv('GOOGLE_API_KEY')
if not API_KEY:
    raise ValueError("GOOGLE_API_KEY environment variable is not set")

chatbot = EnhancedMultiFormatChatbot(API_KEY)
STORE_PER_CHAT_ROWS = False  # set True if you also want per-row storage in messages table

@app.route('/')
# def home():
#     return render_template('index.html')
def index():
    now = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    filename = f"logs/chat_{now}.txt"
    session['chat_log_path'] = filename
    # Ensure a stable client_id is present for this browser session
    _ = get_or_create_client_id()

    full_path = os.path.join(os.getcwd(), filename)
    os.makedirs(os.path.dirname(full_path), exist_ok=True)

    return render_template("login_supabase.html", supabase_url=os.getenv('SUPABASE_URL'), supabase_anon=os.getenv('SUPABASE_ANON_KEY'))

@app.route('/chat')
def chat_supabase():
    return render_template(
        "index.html",
        supabase_url=os.getenv('SUPABASE_URL'),
        supabase_anon=os.getenv('SUPABASE_ANON_KEY')
    )

@app.route('/avatar')
def avatar():
    return render_template("avatar.html")

@app.route('/upload', methods=['POST'])
def upload_files():
    print("[DEBUG] Entered /upload route")
    # Require authenticated user via Supabase
    user_id, auth_err = verify_user(request)
    if auth_err:
        return jsonify(auth_err), 401
    print('[UPLOAD] Authenticated user_id =', user_id)
    # Ensure user profile row exists (if profiles table is present)
    try:
        token = extract_bearer_token(request)
        upsert_profile_as_user(token, user_id, get_user_email(request))
    except Exception:
        pass
    if 'file' not in request.files:
        print("[ERROR] No file part in request")
        return jsonify({'error': 'No file part'}), 400
    
    files = request.files.getlist('file')
    print(f"[DEBUG] Files received: {[file.filename for file in files]}")
    
    if not files or all(file.filename=='' for file in files):
        return jsonify({'error': 'no files selected'}), 400
    

    # Resolve or create client_id
    client_id = request.args.get('client_id') or request.form.get('client_id') or request.headers.get('X-Client-Id') or session.get('client_id') or get_or_create_client_id()
    uploaded_files=[]
    session['uploaded_files']=[]

    combined_text=""
    # success_files=[]
    # failed_files=[]
    
    print("DEBUG allowed_file is:", allowed_file)

    # Load persisted mapping for this client
    client_map = load_client_map()
    current_client_files = client_map.get(client_id, []) if client_id else []

    for file in files:
        if allowed_file(file.filename):  
            filename = str(uuid.uuid4()) + os.path.splitext(file.filename)[1]
            original_name = file.filename or filename
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(filepath)
            print('[UPLOAD] Saved locally:', filepath)

            if 'uploaded_files' not in session:
                session['uploaded_files']=[]
    
            session['uploaded_files'].append(filepath)
            uploaded_files.append(filename)

            if client_id:
                current_client_files.append(filepath)

            # Persist in Supabase for the authenticated user
            try:
                # Upload to Backblaze B2 and store its URL in DB
                object_key = f"users/{user_id}/{filename}"
                print('[UPLOAD] Object key for B2:', object_key)
                storage_url = upload_file_to_b2(filepath, object_key)
                persist_path = storage_url or filepath
                print('[UPLOAD] Persist path stored in DB:', persist_path)
                token = extract_bearer_token(request)
                add_document(token, user_id=user_id, path=persist_path, filename=filename, display_name=original_name)
            except Exception as e:
                print("[DB] add_document failed:", e)

            load_success = chatbot.load_file(filepath)

            combined_text += extract_text_from_pdf(filepath) +"\n\n"

            chatbot.load_combined_text(combined_text)
            session['current_file'] = 'multiple_files_combined'
        else:
            print(f"Skipping unsupported file: {file.filename}")

    # if combined_text:
    #     chatbot.load_combined_text(combined_text)
    #     session['current_file'] = 'multiple_files_combined'

    if uploaded_files:
        # Persist mapping for this client
        if client_id:
            client_map[client_id] = current_client_files
            save_client_map(client_map)
        print('[UPLOAD] client_id mapping updated for', client_id, '->', current_client_files)
        chatbot.load_combined_text(combined_text)
        session['current_file'] = 'multiple_files_combined'
        return jsonify({
            'success': True, 
            'message': f'{len(uploaded_files)} files uploaded and processed successfully', 
            'files': uploaded_files,
            'ready':True,
            'client_id': client_id
            }), 200
    else:
        return jsonify({'error': 'No valid files uploaded'}), 400


@app.route('/ask', methods=['POST'])
def ask_question():
    print("[DEBUG] Reached /ask endpoint")
    data = request.get_json()
    print("[DEBUG] Raw JSON received:", data)

    # Require authenticated user via Supabase
    user_id, auth_err = verify_user(request)
    if auth_err:
        return jsonify(auth_err), 401
    # Ensure user profile row exists (if profiles table is present)
    try:
        token = extract_bearer_token(request)
        upsert_profile_as_user(token, user_id, get_user_email(request))
    except Exception:
        pass

    question = data.get('question')
    mute=data.get('mute',False)
    chat_id = data.get('chat_id')

    # print(f"received : {question} ,mute: {mute}")

    if not question:
        return jsonify({'error': 'No question provided'}), 400
    
    # Resolve client identity
    client_id = (data.get('client_id') if isinstance(data, dict) else None) or session.get('client_id') or get_or_create_client_id()

    # Always load stored document paths from DB for this authenticated user
    try:
        token = extract_bearer_token(request)
        user_docs = list_user_documents(token, user_id)
        print(f"[ASK] DB returned {len(user_docs or [])} document entries for user {user_id}")
        all_files = [doc.get('path') for doc in (user_docs or []) if doc and doc.get('path')]
        # Deduplicate paths while preserving order
        seen_paths = set()
        deduped = []
        for p in all_files:
            if p not in seen_paths:
                seen_paths.add(p)
                deduped.append(p)
        all_files = deduped
        print(f"[ASK] Using {len(all_files)} unique document paths from DB")
        for idx, p in enumerate(all_files):
            print(f"[ASK] doc[{idx+1}]: {p}")
    except Exception as e:
        print('[DB] list_user_documents failed:', e)
        all_files = []

    # Fallback to session/persist map only if DB has no records (first-time users)
    if not all_files:
        session_files = session.get('uploaded_files', [])
        persisted_files = []
        if client_id:
            client_map = load_client_map()
            persisted_files = client_map.get(client_id, [])
        seen = set()
        for p in (session_files or []) + (persisted_files or []):
            if p and p not in seen:
                all_files.append(p)
                seen.add(p)
        if not all_files:
            return jsonify({'error': 'Please upload a file first'}), 400

    # Build combined context
    combined_text = ""
    for filepath in all_files:
        try:
            print('[ASK] Reading from path:', filepath)
            local_path = ensure_local_file(filepath)
            if local_path != filepath:
                print('[ASK] Downloaded remote file to temp path:', local_path)
            # Try pdfplumber first
            extracted = extract_text_from_pdf(local_path)
            if not extracted or len(extracted.strip()) == 0:
                print('[ASK] pdfplumber yielded no text, trying PyPDF2 fallback...')
                try:
                    with open(local_path, 'rb') as f:
                        reader = PdfReader(f)
                        buf = []
                        for pg in reader.pages:
                            buf.append(pg.extract_text() or '')
                        extracted = '\n'.join(buf)
                except Exception as e2:
                    print('[ASK] PyPDF2 fallback failed:', e2)
                    extracted = ''
            if extracted:
                print(f"[ASK] Extracted {len(extracted)} chars from: {os.path.basename(local_path)}")
                combined_text += extracted + "\n\n"
            else:
                print(f"[ASK] No text extracted from: {os.path.basename(local_path)}")
        except Exception as e:
            print(f"[ERROR] Failed to extract text from {filepath}: {e}")
    print(f"[ASK] Total combined context size: {len(combined_text)} chars from {len(all_files)} docs")
    # Initialize chat context for this request and keep session in sync
    try:
        chatbot.load_combined_text(combined_text)
        session['uploaded_files'] = all_files
        session['current_file'] = 'multiple_files_combined'
    except Exception as e:
        print(f"[ERROR] Failed to load combined context into model: {e}")

    try:
        # Ensure chat exists and store messages to DB
        if not chat_id and STORE_PER_CHAT_ROWS:
            try:
                # Create a chat if none provided
                new_chat = create_chat(user_id=user_id, title=f"Chat {datetime.now().strftime('%Y-%m-%d %H:%M')}")
                chat_id = new_chat.get('id') if isinstance(new_chat, dict) else None
            except Exception as ce:
                print('[DB] create_chat failed:', ce)

        # Save user message (both per-chat table and per-user aggregated history)
        try:
            token = extract_bearer_token(request)
            if STORE_PER_CHAT_ROWS and chat_id:
                add_message(chat_id=chat_id, role='user', content=question)
            append_user_message_as_user(token, user_id=user_id, role='user', content=question)
        except Exception as me:
            print('[DB] add_message(user) failed:', me)

        response_text = chatbot.ask_question(question)
        # Final fallback: if user asked for a specific leave type, force concise days-only answer
        try:
            from content_filter import detect_leave_type_from_question, extract_leave_days
            detected_type = detect_leave_type_from_question(question)
            if detected_type:
                # Prefer latest-year filtered content available in chatbot, else combined_text, else model output
                preferred_source = getattr(chatbot, 'filtered_content', '') or combined_text or ''
                concise = extract_leave_days(preferred_source, detected_type) or extract_leave_days(response_text, detected_type)
                if concise:
                    response_text = concise
        except Exception:
            pass
        print("[DEBUG] Got response:", response_text)

        log_conversation(question, response_text)

        if not mute:
            audio_url= generate_polly_audio(response_text)
        else:
            audio_url=None
        # Save assistant message (both per-chat table and per-user aggregated history)
        try:
            token = extract_bearer_token(request)
            if STORE_PER_CHAT_ROWS and chat_id and isinstance(response_text, str):
                add_message(chat_id=chat_id, role='assistant', content=response_text)
            if isinstance(response_text, str):
                append_user_message_as_user(token, user_id=user_id, role='assistant', content=response_text)
        except Exception as me:
            print('[DB] add_message(assistant) failed:', me)

        return jsonify({
            'response': response_text,
            'audio_url': audio_url,
            'client_id': client_id,
            'chat_id': chat_id
            })
    
    except Exception as e:
        print("[ERROR] Exception occurred while generating response:", str(e))
        return jsonify({'error': str(e)}), 500

@app.route('/upload_log', methods=['POST'])
def upload_log():
    uploaded_file=request.files.get('logfile')
    override = request.form.get('override') == 'true'
    if uploaded_file:
        mode = 'w' if override else 'a'
        with open(CHAT_LOG_PATH, mode, encoding='utf-8') as f:
            f.write(uploaded_file.read().decode('utf-8'))
        return redirect('/')
    

@app.route('/test-pdf')
def test_pdf():
    text = extract_text_from_pdf('your/test/path/file.pdf')
    return jsonify({'text': text})

@app.route('/static/<path:filename>')
def serve_static(filename):
    return send_from_directory('static', filename)

@app.route('/auth/sync', methods=['POST'])
def auth_sync():
    try:
        user_id, auth_err = verify_user(request)
        if auth_err:
            return jsonify(auth_err), 401
        token = extract_bearer_token(request)
        email = get_user_email(request)
        profile = upsert_profile_as_user(token, user_id, email)
        return jsonify({'ok': True, 'profile': profile or {'id': user_id, 'email': email}})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/documents', methods=['GET'])
def list_documents_endpoint():
    user_id, auth_err = verify_user(request)
    if auth_err:
        return jsonify(auth_err), 401
    try:
        token = extract_bearer_token(request)
        docs = list_user_documents(token, user_id)
        return jsonify({'documents': docs or []})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/users', methods=['GET'])
def list_users_endpoint():
    # Simple admin-protected endpoint using a static token
    admin_token = os.getenv('ADMIN_TOKEN')
    provided = request.headers.get('X-Admin-Token')
    if not admin_token or provided != admin_token:
        return jsonify({'error': 'Unauthorized'}), 401
    try:
        profiles = list_profiles()
        return jsonify({'users': profiles})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/chats', methods=['GET', 'POST'])
def chats_handler():
    user_id, auth_err = verify_user(request)
    if auth_err:
        return jsonify(auth_err), 401
    if request.method == 'GET':
        try:
            chats = list_user_chats(user_id)
            return jsonify({'chats': chats or []})
        except Exception as e:
            return jsonify({'error': str(e)}), 500
    # POST create chat
    try:
        payload = request.get_json(silent=True) or {}
        title = payload.get('title') or f"Chat {datetime.now().strftime('%Y-%m-%d %H:%M')}"
        chat = create_chat(user_id=user_id, title=title)
        return jsonify({'chat': chat}), 200
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/chats/<chat_id>/messages', methods=['GET'])
def chat_messages(chat_id):
    user_id, auth_err = verify_user(request)
    if auth_err:
        return jsonify(auth_err), 401
    try:
        msgs = list_messages(chat_id)
        return jsonify({'messages': msgs or []})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/user/messages', methods=['GET'])
def user_messages_endpoint():
    user_id, auth_err = verify_user(request)
    if auth_err:
        return jsonify(auth_err), 401
    try:
        token = extract_bearer_token(request)
        msgs = get_user_messages_as_user(token, user_id)
        return jsonify({'messages': msgs})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.errorhandler(404)
def not_found_error(error):
    return jsonify({'error': 'Not found'}), 404

@app.errorhandler(500)
def internal_error(error):
    return jsonify({'error': 'Internal server error'}), 500

if __name__ == "__main__":
    # Check required dependencies
    try:
        pytesseract.get_tesseract_version()
        print("✅ OCR support available")
    except:
        print("⚠️  OCR not available - install Tesseract OCR for image support")
        print("Download from: https://github.com/tesseract-ocr/tesseract")
    
    try:
        import cv2
        print("✅ Enhanced table detection available")
    except ImportError:
        print("⚠️  OpenCV not installed - enhanced table detection not available")
        print("Install with: pip install opencv-python")
    
    # Run the app
    port = int(os.getenv('PORT', 5002))
    app.run(host='0.0.0.0', port=port, debug=True)
    